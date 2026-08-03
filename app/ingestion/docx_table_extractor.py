from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class ExtractedTable:
    table_index: int
    row_count: int
    column_count: int
    rows: list[list[str]]
    text_representation: str
    preceding_paragraph_index: int | None
    preceding_paragraph_text: str | None


@dataclass(frozen=True)
class ExtractedDocxTables:
    file_name: str
    table_count: int
    tables: list[ExtractedTable]


def _normalize_cell_text(text: str) -> str:
    return " ".join(text.split())


def extract_docx_tables(file_path: str | Path) -> ExtractedDocxTables:
    """Extract table content from a DOCX file.

    This step handles tables separately from paragraph extraction so table behavior
    can be tested and debugged independently.
    """

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file does not exist: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected a .docx file, got: {path.name}")

    document = Document(path)
    parent_context_by_table_index = _extract_parent_context_by_table_index(document)
    extracted_tables: list[ExtractedTable] = []

    for table_index, table in enumerate(document.tables):
        rows: list[list[str]] = []
        max_columns = 0

        for row in table.rows:
            normalized_row = [_normalize_cell_text(cell.text) for cell in row.cells]
            rows.append(normalized_row)
            max_columns = max(max_columns, len(normalized_row))

        table_text = "\n".join(" | ".join(row) for row in rows)

        extracted_tables.append(
            ExtractedTable(
                table_index=table_index,
                row_count=len(rows),
                column_count=max_columns,
                rows=rows,
                text_representation=table_text,
                preceding_paragraph_index=parent_context_by_table_index[table_index][0],
                preceding_paragraph_text=parent_context_by_table_index[table_index][1],
            )
        )

    return ExtractedDocxTables(
        file_name=path.name,
        table_count=len(extracted_tables),
        tables=extracted_tables,
    )


def _extract_parent_context_by_table_index(
    document: Document,
) -> list[tuple[int | None, str | None]]:
    """Return the nearest preceding top-level paragraph for each top-level table."""

    contexts: list[tuple[int | None, str | None]] = []
    preceding_index: int | None = None
    preceding_text: str | None = None
    paragraph_index = 0

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = _normalize_cell_text(Paragraph(child, document).text)
            if text:
                preceding_index = paragraph_index
                preceding_text = text
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            contexts.append((preceding_index, preceding_text))

    if len(contexts) != len(document.tables):
        raise RuntimeError(
            "Unable to establish stable top-level DOCX table order for parent context extraction."
        )
    return contexts
