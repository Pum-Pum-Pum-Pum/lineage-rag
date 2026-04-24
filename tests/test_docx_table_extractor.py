from pathlib import Path

from docx import Document

from app.ingestion.docx_table_extractor import extract_docx_tables


def test_extract_docx_tables_basic_case(tmp_path: Path) -> None:
    file_path = tmp_path / "sample_tables.docx"

    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Branch Code"
    table.cell(1, 1).text = "Unique branch identifier"
    document.save(file_path)

    extracted = extract_docx_tables(file_path)

    assert extracted.file_name == "sample_tables.docx"
    assert extracted.table_count == 1
    assert extracted.tables[0].row_count == 2
    assert extracted.tables[0].column_count == 2
    assert extracted.tables[0].rows == [
        ["Column", "Description"],
        ["Branch Code", "Unique branch identifier"],
    ]
    assert extracted.tables[0].text_representation == (
        "Column | Description\n"
        "Branch Code | Unique branch identifier"
    )


def test_extract_docx_tables_missing_file() -> None:
    try:
        extract_docx_tables("missing_tables.docx")
    except FileNotFoundError as exc:
        assert "DOCX file does not exist" in str(exc)
    else:
        raise AssertionError("Expected FileNotFoundError for missing DOCX file")


def test_extract_docx_tables_invalid_suffix(tmp_path: Path) -> None:
    text_file = tmp_path / "not_a_doc.txt"
    text_file.write_text("hello")

    try:
        extract_docx_tables(text_file)
    except ValueError as exc:
        assert "Expected a .docx file" in str(exc)
    else:
        raise AssertionError("Expected ValueError for non-DOCX input")
