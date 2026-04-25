from pathlib import Path

from docx import Document

from app.ingestion.chunker import chunk_normalized_artifact
from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.normalized_artifact import build_normalized_artifact
from app.ingestion.retrieval_ready_artifact import build_retrieval_ready_artifact
from app.ingestion.table_chunker import chunk_tables_from_artifact


def test_build_retrieval_ready_artifact_combines_paragraph_and_table_units(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Paragraph 1")
    document.add_paragraph("Paragraph 2")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Branch Code"
    table.cell(1, 1).text = "Unique branch identifier"
    document.save(file_path)

    raw_artifact = ingest_docx_file(file_path)
    normalized_artifact = build_normalized_artifact(raw_artifact)
    paragraph_chunks = chunk_normalized_artifact(normalized_artifact, max_paragraphs_per_chunk=1)
    table_chunks = chunk_tables_from_artifact(normalized_artifact)

    retrieval_ready = build_retrieval_ready_artifact(
        normalized_artifact,
        paragraph_chunks,
        table_chunks,
    )

    assert retrieval_ready.document_family == "FS_FCIS_14.7.0.0.0$ASNB"
    assert retrieval_ready.release_label == "R24"
    assert retrieval_ready.total_units == 3
    assert retrieval_ready.units[0].source_kind == "paragraph"
    assert retrieval_ready.units[1].source_kind == "paragraph"
    assert retrieval_ready.units[2].source_kind == "table"
    assert retrieval_ready.units[2].text == (
        "Column | Description\n"
        "Branch Code | Unique branch identifier"
    )
