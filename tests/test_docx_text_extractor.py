from pathlib import Path

from docx import Document

from app.ingestion.docx_text_extractor import extract_docx_text


def test_extract_docx_text_paragraphs(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"

    document = Document()
    document.add_paragraph("Functional Overview")
    document.add_paragraph("")
    document.add_paragraph("User can access branch reports.")
    document.save(file_path)

    extracted = extract_docx_text(file_path)

    assert extracted.file_name == "sample.docx"
    assert extracted.paragraph_count == 3
    assert extracted.non_empty_paragraph_count == 2
    assert extracted.paragraphs == [
        "Functional Overview",
        "",
        "User can access branch reports.",
    ]
    assert extracted.full_text == "Functional Overview\nUser can access branch reports."


def test_extract_docx_text_missing_file() -> None:
    try:
        extract_docx_text("missing.docx")
    except FileNotFoundError as exc:
        assert "DOCX file does not exist" in str(exc)
    else:
        raise AssertionError("Expected FileNotFoundError for missing DOCX file")


def test_extract_docx_text_invalid_suffix(tmp_path: Path) -> None:
    text_file = tmp_path / "not_a_doc.txt"
    text_file.write_text("hello")

    try:
        extract_docx_text(text_file)
    except ValueError as exc:
        assert "Expected a .docx file" in str(exc)
    else:
        raise AssertionError("Expected ValueError for non-DOCX input")
