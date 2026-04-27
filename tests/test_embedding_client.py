from pathlib import Path

from docx import Document

from app.embeddings.client import embed_batch
from app.embeddings.embedding_contract import EmbeddingBatch, build_embedding_batch_contract
from app.ingestion.chunker import chunk_normalized_artifact
from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.normalized_artifact import build_normalized_artifact
from app.ingestion.retrieval_ready_artifact import build_retrieval_ready_artifact
from app.ingestion.table_chunker import chunk_tables_from_artifact


class FakeEmbeddingItem:
    def __init__(self, embedding: list[float]) -> None:
        self.embedding = embedding


class FakeEmbeddingResponse:
    def __init__(self, data: list[FakeEmbeddingItem]) -> None:
        self.data = data


class FakeEmbeddingsAPI:
    def create(self, model: str, input: list[str]) -> FakeEmbeddingResponse:
        return FakeEmbeddingResponse(
            [FakeEmbeddingItem([float(index), float(index) + 0.5]) for index, _ in enumerate(input)]
        )


class FakeOpenAIClient:
    def __init__(self) -> None:
        self.embeddings = FakeEmbeddingsAPI()


class MismatchedEmbeddingsAPI:
    def create(self, model: str, input: list[str]) -> FakeEmbeddingResponse:
        return FakeEmbeddingResponse([FakeEmbeddingItem([1.0, 2.0])])


class MismatchedFakeOpenAIClient:
    def __init__(self) -> None:
        self.embeddings = MismatchedEmbeddingsAPI()


def test_embed_batch_updates_status_and_vectors(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Paragraph 1")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Branch Code"
    table.cell(1, 1).text = "Unique branch identifier"
    document.save(file_path)

    raw_artifact = ingest_docx_file(file_path)
    normalized_artifact = build_normalized_artifact(raw_artifact)
    paragraph_chunks = chunk_normalized_artifact(normalized_artifact, max_paragraphs_per_chunk=1)
    table_chunks = chunk_tables_from_artifact(normalized_artifact)
    retrieval_ready = build_retrieval_ready_artifact(
        normalized_artifact,
        paragraph_chunks,
        table_chunks,
    )
    batch = build_embedding_batch_contract(retrieval_ready, embedding_model="text-embedding-3-large")

    embedded_batch = embed_batch(batch, client=FakeOpenAIClient())

    assert embedded_batch.total_records == batch.total_records
    assert embedded_batch.records[0].embedding_status == "embedded"
    assert embedded_batch.records[0].vector == [0.0, 0.5]
    assert embedded_batch.records[1].vector == [1.0, 1.5]


def test_embed_batch_returns_empty_batch_without_api_call() -> None:
    empty_batch = EmbeddingBatch(
        document_name="empty.docx",
        total_records=0,
        records=[],
    )

    embedded_batch = embed_batch(empty_batch, client=FakeOpenAIClient())

    assert embedded_batch.document_name == "empty.docx"
    assert embedded_batch.total_records == 0
    assert embedded_batch.records == []


def test_embed_batch_raises_when_response_count_mismatches_input(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Paragraph 1")
    document.add_paragraph("Paragraph 2")
    document.save(file_path)

    raw_artifact = ingest_docx_file(file_path)
    normalized_artifact = build_normalized_artifact(raw_artifact)
    paragraph_chunks = chunk_normalized_artifact(normalized_artifact, max_paragraphs_per_chunk=1)
    table_chunks = chunk_tables_from_artifact(normalized_artifact)
    retrieval_ready = build_retrieval_ready_artifact(
        normalized_artifact,
        paragraph_chunks,
        table_chunks,
    )
    batch = build_embedding_batch_contract(retrieval_ready, embedding_model="text-embedding-3-large")

    try:
        embed_batch(batch, client=MismatchedFakeOpenAIClient())
    except RuntimeError as exc:
        assert "Embedding response count does not match input record count" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError for mismatched embedding response count")
