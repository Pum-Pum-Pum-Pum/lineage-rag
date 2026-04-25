import json
from pathlib import Path

from docx import Document

from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.normalized_artifact import build_normalized_artifact
from app.ingestion.chunker import chunk_normalized_artifact
from app.ingestion.chunked_artifact_writer import write_chunked_document_to_json


def test_write_chunked_document_to_json(tmp_path: Path) -> None:
    input_file = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Paragraph 1")
    document.add_paragraph("Paragraph 2")
    document.add_paragraph("Paragraph 3")
    document.save(input_file)

    raw_artifact = ingest_docx_file(input_file)
    normalized_artifact = build_normalized_artifact(raw_artifact)
    chunked_document = chunk_normalized_artifact(normalized_artifact, max_paragraphs_per_chunk=2)
    output_file = write_chunked_document_to_json(chunked_document, tmp_path / "processed")

    assert output_file.exists()
    assert output_file.name == "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.chunks.json"

    payload = json.loads(output_file.read_text(encoding="utf-8"))
    assert payload["document_family"] == "FS_FCIS_14.7.0.0.0$ASNB"
    assert payload["release_label"] == "R24"
    assert payload["total_chunks"] == 2
    assert payload["chunks"][0]["text"] == "Paragraph 1\nParagraph 2"
