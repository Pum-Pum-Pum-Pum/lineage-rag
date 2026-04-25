import json
from pathlib import Path

from docx import Document

from app.ingestion.chunker import chunk_normalized_artifact
from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.normalized_artifact import build_normalized_artifact
from app.ingestion.retrieval_ready_artifact import build_retrieval_ready_artifact
from app.ingestion.retrieval_ready_artifact_writer import write_retrieval_ready_artifact_to_json
from app.ingestion.table_chunker import chunk_tables_from_artifact


def test_write_retrieval_ready_artifact_to_json(tmp_path: Path) -> None:
    input_file = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Paragraph 1")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Branch Code"
    table.cell(1, 1).text = "Unique branch identifier"
    document.save(input_file)

    raw_artifact = ingest_docx_file(input_file)
    normalized_artifact = build_normalized_artifact(raw_artifact)
    paragraph_chunks = chunk_normalized_artifact(normalized_artifact, max_paragraphs_per_chunk=1)
    table_chunks = chunk_tables_from_artifact(normalized_artifact)
    retrieval_ready = build_retrieval_ready_artifact(
        normalized_artifact,
        paragraph_chunks,
        table_chunks,
    )

    output_file = write_retrieval_ready_artifact_to_json(retrieval_ready, tmp_path / "processed")

    assert output_file.exists()
    assert output_file.name == "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.retrieval_ready.json"

    payload = json.loads(output_file.read_text(encoding="utf-8"))
    assert payload["document_family"] == "FS_FCIS_14.7.0.0.0$ASNB"
    assert payload["release_label"] == "R24"
    assert payload["total_units"] == 2
    assert payload["units"][0]["source_kind"] == "paragraph"
    assert payload["units"][1]["source_kind"] == "table"
