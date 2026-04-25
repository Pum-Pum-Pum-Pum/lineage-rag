import json
from pathlib import Path

from docx import Document

from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.processed_artifact_writer import write_ingested_artifact_to_json


def test_write_ingested_artifact_to_json(tmp_path: Path) -> None:
    input_file = tmp_path / "FS_FCIS_14.4.0.0.0$ASNB_R2_PNB_Branch Online Reports(BOR)_v1.2.docx"

    document = Document()
    document.add_paragraph("Functional Overview")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Branch Code"
    table.cell(1, 1).text = "Unique branch identifier"
    document.save(input_file)

    artifact = ingest_docx_file(input_file)
    output_file = write_ingested_artifact_to_json(artifact, tmp_path / "processed")

    assert output_file.exists()
    assert output_file.name == "FS_FCIS_14.4.0.0.0$ASNB_R2_PNB_Branch Online Reports(BOR)_v1.2.json"

    payload = json.loads(output_file.read_text(encoding="utf-8"))
    assert payload["document_name"] == input_file.name
    assert payload["parsed_name"]["release_label"] == "R2"
    assert payload["extracted_text"]["non_empty_paragraph_count"] == 1
    assert payload["extracted_tables"]["table_count"] == 1
