from pathlib import Path

from docx import Document

from app.ingestion.docx_ingestion_artifact import ingest_docx_file


def test_ingest_docx_file_combines_name_text_and_tables(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.4.0.0.0$ASNB_R2_PNB_Branch Online Reports(BOR)_v1.2.docx"

    document = Document()
    document.add_paragraph("Functional Overview")
    document.add_paragraph("User can access branch reports.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Branch Code"
    table.cell(1, 1).text = "Unique branch identifier"
    document.save(file_path)

    artifact = ingest_docx_file(file_path)

    assert artifact.document_name == file_path.name
    assert artifact.parsed_name.document_family == "FS_FCIS_14.4.0.0.0$ASNB"
    assert artifact.parsed_name.release_label == "R2"
    assert artifact.parsed_name.variant_suffix == "PNB_Branch Online Reports(BOR)_v1.2"
    assert artifact.extracted_text.non_empty_paragraph_count == 2
    assert artifact.extracted_tables.table_count == 1
    assert artifact.extracted_tables.tables[0].rows[1][0] == "Branch Code"
