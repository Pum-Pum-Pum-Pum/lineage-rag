from pathlib import Path

from docx import Document

from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.normalized_artifact import build_normalized_artifact


def test_build_normalized_artifact_keeps_raw_and_cleaned_views(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Functional Design Document")
    document.add_paragraph(
        "Title, Subject, Last Updated Date, Reference Number, and Version are marked by a Word Bookmark so that they can be easily reproduced in the header and footer of documents."
    )
    document.add_paragraph("Requirements Overview")
    document.save(file_path)

    raw_artifact = ingest_docx_file(file_path)
    normalized_artifact = build_normalized_artifact(raw_artifact)

    assert normalized_artifact.raw_artifact.document_name == file_path.name
    assert normalized_artifact.raw_artifact.extracted_text.non_empty_paragraph_count == 3
    assert normalized_artifact.normalized_text.cleaned_paragraph_count == 2
    assert normalized_artifact.normalized_text.cleaned_paragraphs == [
        "Functional Design Document",
        "Requirements Overview",
    ]
