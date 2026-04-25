from pathlib import Path

from docx import Document

from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.text_normalizer import normalize_ingested_text


def test_normalize_ingested_text_removes_known_noise(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Functional Design Document")
    document.add_paragraph("")
    document.add_paragraph(
        "Title, Subject, Last Updated Date, Reference Number, and Version are marked by a Word Bookmark so that they can be easily reproduced in the header and footer of documents."
    )
    document.add_paragraph("Requirements Overview")
    document.add_paragraph("To add additional approval lines, press [Tab] from the last cell in the table above.")
    document.save(file_path)

    artifact = ingest_docx_file(file_path)
    normalized = normalize_ingested_text(artifact)

    assert normalized.original_non_empty_paragraph_count == 4
    assert normalized.cleaned_paragraph_count == 2
    assert normalized.removed_paragraph_count == 2
    assert normalized.cleaned_paragraphs == [
        "Functional Design Document",
        "Requirements Overview",
    ]
    assert normalized.cleaned_full_text == "Functional Design Document\nRequirements Overview"


def test_normalize_ingested_text_removes_toc_and_front_matter_noise(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Document Control")
    document.add_paragraph("Table of Contents")
    document.add_paragraph("1\tIntroduction\t9")
    document.add_paragraph("1.1\tBackground\t9")
    document.add_paragraph("Functional Overview")
    document.add_paragraph("Requirements Summary")
    document.save(file_path)

    artifact = ingest_docx_file(file_path)
    normalized = normalize_ingested_text(artifact)

    assert normalized.original_non_empty_paragraph_count == 6
    assert normalized.cleaned_paragraph_count == 2
    assert normalized.cleaned_paragraphs == [
        "Functional Overview",
        "Requirements Summary",
    ]


def test_normalize_ingested_text_removes_metadata_and_footer_noise(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Client Name                                   : PNB")
    document.add_paragraph("Release Name                               : FCIS_14.7.0.0.0$ASNB_R24")
    document.add_paragraph("Functional Overview")
    document.add_paragraph("World Headquarters")
    document.add_paragraph("Phone: +1.650.506.7000")
    document.add_paragraph("www.oracle.com/ financial_services/")
    document.add_paragraph("Requirements Summary")
    document.save(file_path)

    artifact = ingest_docx_file(file_path)
    normalized = normalize_ingested_text(artifact)

    assert normalized.cleaned_paragraphs == [
        "Functional Overview",
        "Requirements Summary",
    ]


def test_normalize_ingested_text_removes_oracle_legal_footer_noise(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Functional Overview")
    document.add_paragraph("Oracle Corporation")
    document.add_paragraph("Copyright © 2025. All rights reserved.")
    document.add_paragraph(
        "No part of this work may be reproduced, stored in a retrieval system, adopted or transmitted in any form or by any means."
    )
    document.add_paragraph(
        "Due care has been taken to make this document as accurate as possible. However, Oracle Financial Services Software Limited makes no representation or warranties with respect to the contents hereof."
    )
    document.add_paragraph(
        "All company and product names are trademarks of the respective companies with which they are associated."
    )
    document.add_paragraph("Requirements Summary")
    document.save(file_path)

    artifact = ingest_docx_file(file_path)
    normalized = normalize_ingested_text(artifact)

    assert normalized.cleaned_paragraphs == [
        "Functional Overview",
        "Requirements Summary",
    ]
