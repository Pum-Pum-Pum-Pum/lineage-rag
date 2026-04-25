from pathlib import Path

from docx import Document

from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.normalized_artifact import build_normalized_artifact
from app.ingestion.chunker import chunk_normalized_artifact


def test_chunk_normalized_artifact_basic_case(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Paragraph 1")
    document.add_paragraph("Paragraph 2")
    document.add_paragraph("Paragraph 3")
    document.add_paragraph("Paragraph 4")
    document.add_paragraph("Paragraph 5")
    document.add_paragraph("Paragraph 6")
    document.save(file_path)

    raw_artifact = ingest_docx_file(file_path)
    normalized_artifact = build_normalized_artifact(raw_artifact)
    chunked = chunk_normalized_artifact(normalized_artifact, max_paragraphs_per_chunk=2)

    assert chunked.document_family == "FS_FCIS_14.7.0.0.0$ASNB"
    assert chunked.release_label == "R24"
    assert chunked.total_chunks == 3
    assert chunked.chunks[0].paragraph_count == 2
    assert chunked.chunks[0].text == "Paragraph 1\nParagraph 2"
    assert chunked.chunks[2].text == "Paragraph 5\nParagraph 6"


def test_chunk_normalized_artifact_invalid_chunk_size(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Paragraph 1")
    document.save(file_path)

    raw_artifact = ingest_docx_file(file_path)
    normalized_artifact = build_normalized_artifact(raw_artifact)

    try:
        chunk_normalized_artifact(normalized_artifact, max_paragraphs_per_chunk=0)
    except ValueError as exc:
        assert "must be greater than 0" in str(exc)
    else:
        raise AssertionError("Expected ValueError for invalid chunk size")
