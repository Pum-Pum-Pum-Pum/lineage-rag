from pathlib import Path

from docx import Document

from app.ingestion.chunker import chunk_normalized_artifact
from app.ingestion.docx_ingestion_artifact import ingest_docx_file
from app.ingestion.normalized_artifact import build_normalized_artifact
from app.ingestion.retrieval_ready_artifact import build_retrieval_ready_artifact
from app.ingestion.table_chunker import chunk_tables_from_artifact
from app.embeddings.embedding_contract import (
    build_embedding_batch_contract,
    compute_content_hash,
    compute_embedding_cache_key,
)


def test_build_embedding_batch_contract(tmp_path: Path) -> None:
    file_path = tmp_path / "FS_FCIS_14.7.0.0.0$ASNB_R24_Teller_Branch_Reports_Realignment_v1.0.docx"

    document = Document()
    document.add_paragraph("Paragraph 1")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "Branch Code"
    table.cell(1, 1).text = "Unique branch identifier"
    document.save(file_path)

    raw_artifact = ingest_docx_file(file_path)
    normalized_artifact = build_normalized_artifact(raw_artifact)
    paragraph_chunks = chunk_normalized_artifact(normalized_artifact, max_paragraphs_per_chunk=1)
    table_chunks = chunk_tables_from_artifact(normalized_artifact)
    retrieval_ready = build_retrieval_ready_artifact(
        normalized_artifact,
        paragraph_chunks,
        table_chunks,
    )

    batch = build_embedding_batch_contract(
        retrieval_ready,
        embedding_model="text-embedding-3-large",
        artifact_version="v1",
    )

    assert batch.document_name == file_path.name
    assert batch.total_records == 2
    assert batch.records[0].content_hash == compute_content_hash(batch.records[0].text)
    assert batch.records[0].artifact_version == "v1"
    assert batch.records[0].cache_key == compute_embedding_cache_key(
        content_hash=batch.records[0].content_hash,
        embedding_model="text-embedding-3-large",
        artifact_version="v1",
    )
    assert batch.records[0].embedding_model == "text-embedding-3-large"
    assert batch.records[0].embedding_status == "pending"
    assert batch.records[0].vector is None
    assert batch.records[1].source_kind == "table"


def test_compute_content_hash_normalizes_line_endings() -> None:
    assert compute_content_hash("line 1\r\nline 2") == compute_content_hash("line 1\nline 2")


def test_compute_embedding_cache_key_changes_when_model_or_artifact_version_changes() -> None:
    content_hash = compute_content_hash("same text")

    baseline_key = compute_embedding_cache_key(
        content_hash=content_hash,
        embedding_model="text-embedding-3-large",
        artifact_version="v1",
    )
    changed_model_key = compute_embedding_cache_key(
        content_hash=content_hash,
        embedding_model="text-embedding-3-small",
        artifact_version="v1",
    )
    changed_artifact_key = compute_embedding_cache_key(
        content_hash=content_hash,
        embedding_model="text-embedding-3-large",
        artifact_version="v2",
    )

    assert baseline_key != changed_model_key
    assert baseline_key != changed_artifact_key
